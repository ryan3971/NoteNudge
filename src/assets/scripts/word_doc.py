import datetime
from docx import Document
import sys
import json
from htmldocx import HtmlToDocx
from docx.shared import Pt
from docx.oxml.ns import qn


def append_text_to_word_document(text, document_path):
    # Get the current date and format it as "Week of [date]"
    # Get the date for the start if the week
    current_date = datetime.datetime.now() - datetime.timedelta(
        days=datetime.datetime.now().weekday()
    )
    current_date = current_date.strftime("%B %d")
    week_heading = f"Week of {current_date}:"

    # Get the current weekday
    current_weekday = datetime.datetime.now().strftime("%A") + ":"

    # Check if the document is open, if so, prompt the user to close it
    # Check if the document exists, if not create it
    try:
        document = Document(document_path)
    except:
        document = Document()
        document.save(document_path)

    # Set the heading styles
    rFonts = document.styles["Heading 1"].element.rPr.rFonts
    rFonts.set(
        qn("w:asciiTheme"), ""
    )  # fancy method necessary to override default heading font
    document.styles["Heading 1"].font.name = "Calibri Light (Headings)"
    document.styles["Heading 1"].font.bold = False
    document.styles["Heading 1"].font.size = Pt(16)

    rFonts = document.styles["Heading 2"].element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "")
    document.styles["Heading 2"].font.name = "Calibri Light (Headings)"
    document.styles["Heading 2"].font.bold = False
    document.styles["Heading 2"].font.size = Pt(13)

    # Find or create the week heading
    week_heading_found = False
    for paragraph in document.paragraphs:
        if paragraph.text == week_heading:
            week_heading_found = True
            break

    if not week_heading_found:
        # set the heading style to Calibri Light (Headings)
        document.add_heading(week_heading, level=1)

    # Find or create the weekday heading
    weekday_heading_found = False
    for paragraph in document.paragraphs:
        if paragraph.text == current_weekday:
            weekday_heading_found = True
            break

    if not weekday_heading_found:
        document.add_heading(current_weekday, level=2)

    # # Find the entry number for the day
    # entry_number = 1
    # for paragraph in document.paragraphs:
    #     if paragraph.text.startswith("entry #"):
    #         entry_number += 1

    # Create the HTML to DOCX parser
    # Add the text as a new paragraph
    #   document.add_paragraph(f"entry #{entry_number}")
    new_parser = HtmlToDocx()
    new_parser.add_html_to_document(text, document)

    # Save and close the document
    document.save(document_path)
    print("Python script executed successfully")


if __name__ == "__main__":
    print("Python script executing")

    # Read all input from stdin
    input_data = sys.stdin.read()

    # Split the input into lines
    input_lines = input_data.splitlines()
    json_data = input_lines[0]
    document_path = input_lines[1]

    try:
        entry_input_json = json.loads(json_data)
        entry_content = entry_input_json.get("content", "")

        append_text_to_word_document(entry_content, document_path)

        # Optional: Return a success message
        print("Python script executed successfully")

    except Exception as e:
        # Print an error message if an exception occurs
        print(json.dumps({"status": "error", "message": str(e)}), file=sys.stderr)
        sys.exit(1)
