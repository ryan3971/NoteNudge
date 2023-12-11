import datetime
from docx import Document
import sys

def append_text_to_word_document(text):
    # Get the current date and format it as "Week of [date]"
    # Get the date for the start if the week
    current_date = datetime.datetime.now() - datetime.timedelta(days=datetime.datetime.now().weekday())
    current_date = current_date.strftime("%B %d")
    week_heading = f"Week of {current_date}"

    # Get the current weekday
    current_weekday = datetime.datetime.now().strftime("%A")

    # Check if the document is open, if so, prompt the user to close it


    # Check if the document exists, if not create it
    try:
        document = Document("word_doc.docx")
    except:
        document = Document()
        document.save("word_doc.docx")

    # Find or create the week heading
    week_heading_found = False
    for paragraph in document.paragraphs:
        if paragraph.text == week_heading:
            week_heading_found = True
            break

    if not week_heading_found:
        document.add_heading(week_heading, level=1)

    # Find or create the weekday heading
    weekday_heading_found = False
    for paragraph in document.paragraphs:
        if paragraph.text == current_weekday:
            weekday_heading_found = True
            break

    if not weekday_heading_found:
        document.add_heading(current_weekday, level=2)

    # Find the entry number for the day
    entry_number = 1
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("entry #"):
            entry_number += 1

    # Add the text as a new paragraph
    document.add_paragraph(f"entry #{entry_number}: {text}")

    # Save and close the document
    document.save("word_doc.docx")
    
# Process the arguments passed to the script and pass it to the function
text_entry = sys.argv[1]
append_text_to_word_document(text_entry)


# # Create the eetry point to the script. 
# if __name__ == "__main__":
#     text_entry = "Hello world"#input("Enter your text entry: ")
#     append_text_to_word_document(text_entry)